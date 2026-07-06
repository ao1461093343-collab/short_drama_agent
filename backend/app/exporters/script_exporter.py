import tempfile
from pathlib import Path
from uuid import uuid4

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


EXPORT_DIR = Path(tempfile.gettempdir()) / "short_drama_agent_exports"
EXPORT_DIR.mkdir(exist_ok=True)


def export_script_docx(result: dict) -> Path:
    path = EXPORT_DIR / f"short_drama_{uuid4().hex}.docx"
    doc = Document()
    script = result.get("final_script", {})
    bible = result.get("project_bible", {})

    doc.add_heading(script.get("title", "短剧剧本"), level=1)
    doc.add_heading("项目圣经", level=2)
    doc.add_paragraph(bible.get("logline", ""))
    doc.add_paragraph(bible.get("theme", ""))

    doc.add_heading("单集剧本", level=2)
    doc.add_paragraph(f"前三秒钩子：{script.get('hook_3s', '')}")
    for scene in script.get("scenes", []):
        doc.add_heading(f"场景 {scene.get('scene_no')} · {scene.get('location')}", level=3)
        doc.add_paragraph(scene.get("action", ""))
        for line in scene.get("dialogue", []):
            doc.add_paragraph(f"{line.get('speaker')}：{line.get('line')}")

    doc.add_heading("拍摄分镜", level=2)
    table = doc.add_table(rows=1, cols=4)
    headers = ["镜号", "场景", "画面", "台词"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for shot in result.get("shooting_script", []):
        row = table.add_row().cells
        row[0].text = str(shot.get("镜号", ""))
        row[1].text = str(shot.get("场景", ""))
        row[2].text = str(shot.get("画面", ""))
        row[3].text = str(shot.get("台词", ""))

    doc.save(path)
    return path


def export_script_pdf(result: dict) -> Path:
    path = EXPORT_DIR / f"short_drama_{uuid4().hex}.pdf"
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = "STSong-Light"
    body_style = ParagraphStyle(
        "CJKBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=9,
        leading=13,
        wordWrap="CJK",
        splitLongWords=True,
    )
    table_style = ParagraphStyle(
        "CJKTable",
        parent=body_style,
        fontSize=7.5,
        leading=10,
        wordWrap="CJK",
        splitLongWords=True,
    )
    header_style = ParagraphStyle(
        "CJKTableHeader",
        parent=table_style,
        fontSize=8,
        leading=10,
        alignment=1,
    )

    script = result.get("final_script", {})
    bible = result.get("project_bible", {})
    story = [
        Paragraph(script.get("title", "短剧剧本"), styles["Title"]),
        Spacer(1, 12),
        Paragraph("项目圣经", styles["Heading2"]),
        Paragraph(_pdf_text(bible.get("logline", "")), body_style),
        Paragraph(_pdf_text(bible.get("theme", "")), body_style),
        Spacer(1, 12),
        Paragraph("单集剧本", styles["Heading2"]),
        Paragraph(_pdf_text(f"前三秒钩子：{script.get('hook_3s', '')}"), body_style),
    ]

    for scene in script.get("scenes", []):
        story.append(Paragraph(_pdf_text(f"场景 {scene.get('scene_no')} · {scene.get('location')}"), styles["Heading3"]))
        story.append(Paragraph(_pdf_text(scene.get("action", "")), body_style))
        for line in scene.get("dialogue", []):
            story.append(Paragraph(_pdf_text(f"{line.get('speaker')}：{line.get('line')}"), body_style))

    rows = [[
        Paragraph("镜号", header_style),
        Paragraph("场景", header_style),
        Paragraph("画面", header_style),
        Paragraph("台词", header_style),
    ]]
    for shot in result.get("shooting_script", []):
        rows.append([
            Paragraph(_pdf_text(_shot_value(shot, "镜号", "闀滃彿", "shot_no")), table_style),
            Paragraph(_pdf_text(_shot_value(shot, "场景", "鍦烘櫙", "scene")), table_style),
            Paragraph(_pdf_text(_shot_value(shot, "画面", "鐢婚潰", "visual")), table_style),
            Paragraph(_pdf_text(_shot_value(shot, "台词", "鍙拌瘝", "dialogue")), table_style),
        ])
    table = Table(
        rows,
        colWidths=[20 * mm, 34 * mm, 70 * mm, 54 * mm],
        repeatRows=1,
        splitByRow=1,
    )
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([Spacer(1, 12), Paragraph("拍摄分镜", styles["Heading2"]), table])

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )
    doc.build(story)
    return path


def _pdf_text(value) -> str:
    text = str(value or "")
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def _shot_value(shot: dict, *keys: str) -> str:
    for key in keys:
        value = shot.get(key)
        if value not in {None, ""}:
            return str(value)
    return ""
